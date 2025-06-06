import datetime
import json
import os
import uuid

from django.conf import settings
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from django.http import FileResponse
import io
import requests
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.db import transaction
from django.db.models import Count, Q
from django.shortcuts import render, redirect
from django.http import JsonResponse, HttpResponse
from .models import *
from django.views import View
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator

def clear_all_employee_departments():
    employees = Employee.objects.all()

    for employee in employees:
        employee.department.clear()
    print(f"Xodimlar uchun bo'llimlar tozalandi.")

def get_session_id(request):
    if 'device_id' not in request.session:
        request.session['device_id'] = str(uuid.uuid4())
    return request.session['device_id']

def update_and_save_employee(request):
    token = os.environ.get('HEMIS_TOKEN')
    headers = {
        'Authorization': f'Bearer {token}'
    }
    url = 'https://student.tersu.uz/rest/v1/data/employee-list'

    get_page_count_params = {
        'type': 'all',
        'page': 1,
        'limit': 100
    }

    get_page_count_response = requests.get(url, headers=headers, params=get_page_count_params)
    if get_page_count_response.status_code == 200:
        clear_all_employee_departments()
        page_count = get_page_count_response.json()['data']['pagination']['pageCount']
        employees = Employee.objects.filter(is_archive=False)

        employee_api_ids = []
        employee_database_ids = [x.employee_id_number for x in employees]

        for page in range(1, page_count + 1):
            params = {
                'type': 'all',
                'page': page,
                'limit': 100
            }
            response = requests.get(url, headers=headers, params=params)

            if response.status_code == 200:
                items = response.json()['data']['items']
                for item in items:
                    try:
                        employee_api_ids.append(item['employee_id_number'])
                        gender, _ = Gender.objects.get_or_create(code=item['gender']['code'],
                                                                 defaults={'name': item['gender']['name']})

                        academic_degree, _ = AcademicDegree.objects.get_or_create(
                            code=item['academicDegree']['code'],
                            defaults={'name': item['academicDegree']['name']}
                        )

                        academic_rank, _ = AcademicRank.objects.get_or_create(
                            code=item['academicRank']['code'],
                            defaults={'name': item['academicRank']['name']}
                        )

                        structure_type, _ = StructureType.objects.get_or_create(
                            code=item['department']['structureType']['code'],
                            defaults={'name': item['department']['structureType']['name']}
                        )

                        locality_type, _ = LocalityType.objects.get_or_create(
                            code=item['department']['localityType']['code'],
                            defaults={'name': item['department']['localityType']['name']}
                        )

                        employment_form, _ = EmploymentForm.objects.get_or_create(
                            code=item['employmentForm']['code'],
                            defaults={'name': item['employmentForm']['name']}
                        )

                        employment_staff, _ = EmploymentStaff.objects.get_or_create(
                            code=item['employmentStaff']['code'],
                            defaults={'name': item['employmentStaff']['name']}
                        )

                        staff_position, _ = StaffPosition.objects.get_or_create(
                            code=item['staffPosition']['code'],
                            defaults={'name': item['staffPosition']['name']}
                        )

                        employee_status, _ = EmployeeStatus.objects.get_or_create(
                            code=item['employeeStatus']['code'],
                            defaults={'name': item['employeeStatus']['name']}
                        )

                        employee_type, _ = EmployeeType.objects.get_or_create(
                            code=item['employeeType']['code'],
                            defaults={'name': item['employeeType']['name']}
                        )

                        department, _ = Department.objects.get_or_create(
                            id=item['department']['id'],
                            defaults={
                                'name': item['department']['name'],
                                'code': item['department']['code'],
                                'structure_type': structure_type,
                                'locality_type': locality_type,
                                'parent': item['department']['parent'],
                            }
                        )

                        department_branch, _ = DepartmentBranch.objects.get_or_create(
                            department=department,
                            employee_form=employment_form,
                            employee_staff= employment_staff,
                            employee_position= staff_position,
                            employee_status= employee_status,
                            employee_type=employee_type,
                        )

                        employee, created = Employee.objects.update_or_create(
                            employee_id_number=item['employee_id_number'],
                            defaults={
                                'meta_id': item['meta_id'],
                                'full_name': item['full_name'],
                                'short_name': item['short_name'],
                                'first_name': item['first_name'],
                                'second_name': item['second_name'],
                                'third_name': item['third_name'],
                                'gender': gender,
                                'birth_timestamp': item['birth_date'],
                                'image': item['image'],
                                'year_of_enter': item['year_of_enter'],
                                'specialty': item['specialty'],
                                'academicDegree': academic_degree,
                                'academicRank': academic_rank,
                                'decree_number': item['decree_number'],
                                'contract_date': item['contract_date'],
                                'created_at': item['created_at'],
                                'updated_at': item['updated_at'],
                                'hash': item['hash'],
                                'is_archive': False
                            }
                        )
                        employee.department.add(department_branch)
                    except Exception as e:
                        print(e)
                        return HttpResponse(e)
            else:
                print(response.status_code)
                return redirect('teachers:home')

        employee_api_ids = list(dict.fromkeys(employee_api_ids))

        fired_employees = list(set(employee_database_ids) - set(employee_api_ids))
        if fired_employees:
            for employee in fired_employees:
                try:
                    fired = Employee.objects.get(employee_id_number=employee)
                    fired.is_archive = True
                    fired.save()
                except Exception as e:
                    messages.warning(request, e)

        messages.success(request, 'muvaffaqiyatli saqlandi')
        return redirect('teachers:home')

    else:
        messages.error(request, f"get page count error. status code {get_page_count_response.status_code}")
        return redirect('teachers:home')

def home(request):
    session_id = get_session_id(request)
    selected_employees = SelectedEmployee.objects.filter(status=True).exclude(votes__session_id=session_id)
    if selected_employees:
        return render(request, 'start.html')
    else:
        return redirect('teachers:finish_vote')

def finish_vote(request):

    return render(request, 'quiz-result.html')


def start(request):
    if request.user.is_authenticated:
        return render(request, 'table.html')
    else:
        return redirect('teachers:home')


def get_employee_list(request):
    # Employeelarni voted=True bo'lganlarni pastga tushirish uchun saralash
    employees = SelectedEmployee.objects.all().order_by('voted', 'created_at')
    employee_list = []
    for employee in employees:
        employee_list.append({
            'id': employee.id,
            'full_name': employee.employee.full_name,
            'image': employee.employee.image if employee.employee.image else '/static/default.jpg',
            'voted': employee.voted,
            'status': employee.status,
        })
    return JsonResponse({'employees': employee_list})

@method_decorator(csrf_exempt, name='dispatch')
class UpdateEmployeeStatusView(View):
    def post(self, request):
        try:
            employee_id = request.POST.get('employee_id')
            status = request.POST.get('status') == 'true'
            voted = request.POST.get('voted') == 'true'

            employee = SelectedEmployee.objects.get(id=employee_id)

            if status:
                SelectedEmployee.objects.exclude(id=employee_id).update(status=False)

            employee.status = status
            employee.voted = voted

            if employee.linked_employee:
                employee.linked_employee.status = status
                employee.linked_employee.voted = voted
                employee.linked_employee.save()

            employee.save()

            return JsonResponse({
                'status': 'success',
                'message': 'Status successfully updated',
                'linked_employee_id': employee.linked_employee.id if employee.linked_employee else None
            })
        except SelectedEmployee.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Employee not found'
            }, status=404)
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=500)

    def get(self, request):
        return JsonResponse({
            'status': 'error',
            'message': 'Method not allowed'
        }, status=405)


def start_vote(request):
    session_id = get_session_id(request)
    selected_employees = SelectedEmployee.objects.filter(status=True).exclude(votes__session_id=session_id).order_by('created_at')

    if selected_employees:
        return render(request, 'quiz-start.html', {'selected_employees': selected_employees})
    else:
        return redirect('teachers:finish_vote')


from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.views import View
from django.http import JsonResponse
from .models import SelectedEmployee, Vote
from django.shortcuts import get_object_or_404
import json


@method_decorator(csrf_exempt, name='dispatch')
class VoteView(View):
    def post(self, request):
        data = json.loads(request.body)
        employee_id = data.get('employee_id')  # Bu employee_id_number
        vote = data.get('vote')
        is_last = data.get('is_last')

        session_id = get_session_id(request)  # get_session_id funktsiyasi mavjud deb taxmin qilamiz

        selected_employee = get_object_or_404(SelectedEmployee, employee__employee_id_number=employee_id)

        # Agar ushbu employee uchun allaqachon ovoz berilgan bo‘lsa
        if Vote.objects.filter(employee=selected_employee, session_id=session_id).exists():
            return JsonResponse({'success': False, 'message': "You have already voted."})

        # Bog‘langan employee tekshiruvi
        linked_employee = selected_employee.linked_employee
        if linked_employee:
            linked_vote = Vote.objects.filter(employee=linked_employee, session_id=session_id).first()
            if linked_vote and linked_vote.vote == 'rozi' and vote == 'rozi':
                return JsonResponse({'success': False,
                                     'message': "Linked employee already voted 'Roziman'. You cannot vote 'Roziman' for this employee."})

        # Ovoz berishni yaratish
        Vote.objects.create(employee=selected_employee, vote=vote, session_id=session_id)

        response_data = {
            'success': True,
            'message': "Your vote has been recorded.",
            'employee_id': employee_id,
            'vote': vote
        }
        if is_last:
            response_data['is_last'] = True

        return JsonResponse(response_data)

def dashboard(request):
    return render(request, 'counter.html')


def get_dashboard_data(request):
    votes_count = Vote.objects.values('session_id').distinct().count()

    return JsonResponse({
            'votes_count': votes_count,
        })

def get_result(request):
    if request.user.is_authenticated:
        return render(request, 'final_table.html')
    else:
        return redirect('teachers:home')

def get_result_data(request):
    votes_count = Vote.objects.values('session_id').distinct().count()
    employees = SelectedEmployee.objects.all().order_by('created_at')
    all_true = Vote.objects.filter(vote__contains='rozi').count()
    all_false = Vote.objects.filter(vote__contains='qarshi').count()
    all_neutral = Vote.objects.filter(vote__contains='betaraf').count()
    timer = VotingTimer.objects.get(id=1)

    votes_list = [all_false, all_true, all_neutral]

    employees_result = []
    for employee in employees:
        employee_votes = Vote.objects.filter(employee=employee).count()
        employee_true_votes = Vote.objects.filter(employee=employee, vote__contains='rozi').count()
        employee_false_votes = Vote.objects.filter(employee=employee, vote__contains='qarshi').count()
        employee_neutral_votes = Vote.objects.filter(employee=employee, vote__contains='betaraf').count()
        all_votes = Vote.objects.filter(employee=employee).count()

        percentage = (employee_true_votes / employee_votes) * 100 if employee_votes > 0 else 0

        employees_result.append({
            'name': str(employee.employee.full_name),
            'position': str(employee.which_position),
            'image': str(employee.employee.image),
            'votes': employee_votes,
            'true': employee_true_votes,
            'false': employee_false_votes,
            'neutral': employee_neutral_votes,
            'all_votes': all_votes,
            'percentage': round(percentage, 2),
            'status': employee.status,
            'voted': employee.voted
        })

    return JsonResponse({
        'votes_count': votes_count,
        'employees': employees_result,
        'votes_list': votes_list,
    })


def get_vote_data(request):
    employee_id = request.GET.get('employee_id')
    if employee_id:
        active_employee = SelectedEmployee.objects.filter(id=employee_id, status=True, voted=False).first()
    else:
        active_employee = SelectedEmployee.objects.filter(status=True, voted=False).first()

    if active_employee:
        votes = Vote.objects.filter(employee=active_employee)
        dt = datetime.datetime.fromtimestamp(int(active_employee.employee.birth_timestamp))
        vote_counts = [
            votes.filter(vote='qarshi').count(),
            votes.filter(vote='rozi').count(),
            votes.filter(vote='betaraf').count()
        ]

        # Bog‘langan employee ma'lumotlari
        linked_employee_data = None
        linked_employee = active_employee.linked_employee
        if linked_employee and linked_employee.status and not linked_employee.voted:
            linked_votes = Vote.objects.filter(employee=linked_employee)
            linked_vote_counts = [
                linked_votes.filter(vote='qarshi').count(),
                linked_votes.filter(vote='rozi').count(),
                linked_votes.filter(vote='betaraf').count()
            ]
            linked_dt = datetime.datetime.fromtimestamp(int(linked_employee.employee.birth_timestamp))
            linked_employee_data = {
                'employee_id': linked_employee.id,
                'full_name': linked_employee.employee.full_name,
                'image': linked_employee.employee.image if linked_employee.employee.image else '/static/default.jpg',
                'birth_date': linked_dt.strftime('%d.%m.%Y'),
                'chair': linked_employee.employee.department.first().department.name,
                'position': linked_employee.employee.department.first().employee_position.name,
                'degree': linked_employee.employee.academicDegree.name,
                'which_position': linked_employee.which_position,
                'votes': linked_vote_counts
            }

        return JsonResponse({
            'status': 'success',
            'employee_id': active_employee.id,
            'full_name': active_employee.employee.full_name,
            'image': active_employee.employee.image if active_employee.employee.image else '/static/default.jpg',
            'birth_date': dt.strftime('%d.%m.%Y'),
            'chair': active_employee.employee.department.first().department.name,
            'position': active_employee.employee.department.first().employee_position.name,
            'degree': active_employee.employee.academicDegree.name,
            'which_position': active_employee.which_position,
            'votes': vote_counts,
            'linked_employee': linked_employee_data  # Bog‘langan employee qo‘shildi
        })
    return JsonResponse({
        'status': 'no_active',
        'message': 'Hozirda faol ovoz berish mavjud emas'
    })

@csrf_exempt
def save_timer(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        timer, _ = VotingTimer.objects.get_or_create(id=1)
        timer.seconds = data['seconds']
        timer.is_running = data['is_running']
        timer.save()
        return JsonResponse({'status': 'success'})
    return JsonResponse({'status': 'error'}, status=400)

def get_timer(request):
    timer, _ = VotingTimer.objects.get_or_create(id=1)
    return JsonResponse({
        'seconds': timer.seconds,
        'is_running': str(timer.is_running).lower()
    })




def custom_title(text):
    if not text:
        return text

    result = []
    current_word = []
    i = 0

    while i < len(text):
        if text[i].isspace():
            if current_word:
                result.append(''.join(current_word))
                current_word = []
            result.append(text[i])
        elif text[i] in "'":
            if current_word:
                result.append(''.join(current_word))
                current_word = [text[i]]
            else:
                current_word.append(text[i])
        else:
            current_word.append(
                text[i].lower() if i > 0 and text[i - 1].isalpha() else text[i].upper() if not current_word else text[
                    i].lower())
        i += 1

    if current_word:
        result.append(''.join(current_word))

    return ''.join(result)

def generate_vote_results_docx(request):
    static_dir = settings.STATICFILES_DIRS[0]
    template_path = os.path.join(static_dir, 'template.docx')

    try:
        template = Document(template_path)
    except Exception as e:
        raise Exception(f"Template faylini ochishda xatolik: {e}")

    employees = SelectedEmployee.objects.all().order_by('created_at')
    items = []

    for i, employee in enumerate(employees, 1):
        employee_votes = Vote.objects.filter(employee=employee).count()
        employee_true_votes = Vote.objects.filter(employee=employee, vote='rozi').count()  # 'rozi' ni to'g'ri tekshirish
        employee_false_votes = Vote.objects.filter(employee=employee, vote='qarshi').count()
        employee_neutral_votes = Vote.objects.filter(employee=employee, vote='betaraf').count()
        all_votes = Vote.objects.filter(employee=employee).count()

        true_percentage = (employee_true_votes / all_votes * 100) if all_votes > 0 else 0
        false_percentage = (employee_false_votes / all_votes * 100) if all_votes > 0 else 0
        neutral_percentage = (employee_neutral_votes / all_votes * 100) if all_votes > 0 else 0
        percentage = (employee_true_votes / employee_votes) * 100 if employee_votes > 0 else 0

        item = {
            'number': str(i),
            'position': employee.which_position,
            'candidate': custom_title(employee.employee.full_name),
            'total_votes': str(all_votes),
            'yes': f"{employee_true_votes} ({true_percentage:.1f}%)",
            'no': f"{employee_false_votes} ({false_percentage:.1f}%)",
            'neutral': f"{employee_neutral_votes} ({neutral_percentage:.1f}%)",
            'result': 'O\'tdi' if percentage >= 50 else 'O\'tmadi',
        }
        items.append(item)

    try:
        table = template.tables[0]
    except IndexError:
        raise Exception("Hujjatda jadval topilmadi. Templatni tekshiring.")

    if len(table.rows) > 1:
        raise Exception("Templatdagi jadvalda faqat sarlavha qatori bo‘lishi kerak. Qo'shimcha qatorlarni o'chiring.")

    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

    for item in items:
        new_row = table.add_row()
        new_row.cells[0].text = item['number']
        new_row.cells[1].text = item['position']  # Лавозим
        new_row.cells[2].text = item['candidate']  # Номзодлар
        new_row.cells[3].text = item['total_votes']  # Овоз беришда қатнашганлар сони
        new_row.cells[4].text = item['yes']  # Рози
        new_row.cells[5].text = item['no']  # Қарши
        new_row.cells[6].text = item['neutral']  # Бетараф
        new_row.cells[7].text = item['result']  # Натижа

        # Har bir yangi qatorning hujayralarini formatlash
        for cell in new_row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Vertikal markazlash
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Gorizontal markazlash
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)  # Shrift o‘lchamini 14 punkt qilish

    # Hujjatni xotirada saqlash (buffer)
    buffer = io.BytesIO()
    template.save(buffer)
    buffer.seek(0)

    # Foydalanuvchi uchun yuklab olish
    response = FileResponse(buffer, as_attachment=True, filename='vote_results.docx')
    return response




@csrf_exempt
def reset_all_status(request):
    try:
        updated_count = SelectedEmployee.objects.update(status=False, voted=False)
        return JsonResponse({
            'status': 'success',
            'message': f'{updated_count} ta obyektning status va voted qiymatlari False ga o\'zgartirildi'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)
